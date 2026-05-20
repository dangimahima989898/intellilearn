import uuid
import os
from docx import Document
from database import SessionLocal
from models.note import Note

def run_seed():
    os.makedirs("uploads", exist_ok=True)
    
    # 1. Create the docx file
    doc_path = "uploads/dsa_binary_trees.docx"
    doc = Document()
    doc.add_heading("Binary Trees and Traversals", level=1)
    doc.add_paragraph("A binary tree is a hierarchical data structure in which each node has at most two children, referred to as the left child and the right child.")
    doc.add_paragraph("Key terms to remember: Leaf node is a node with no children. Root node is the topmost node of the tree. Height is the number of edges on the longest path from root to a leaf.")
    doc.add_paragraph("Three standard DFS traversals are Preorder, Inorder, and Postorder. Preorder visits root first, then left subtree, then right subtree. Inorder visits left first, then root, then right. Postorder visits left first, then right, then root.")
    doc.save(doc_path)
    
    # 2. Add DB record
    db = SessionLocal()
    try:
        # Check if already exists
        existing = db.query(Note).filter(Note.title == "Binary Trees Study Guide").first()
        if existing:
            print("Note already seeded.")
            return
            
        note = Note(
            id=uuid.uuid4(),
            title="Binary Trees Study Guide",
            subject_id=uuid.UUID("05906049-0c65-4bde-87c0-96f0135a0e09"), # DSA
            file_url="/uploads/dsa_binary_trees.docx",
            file_type="docx",
            file_size_kb=15,
            uploaded_by=uuid.UUID("c8e6ee7c-9def-43b0-9013-4cc4ba875ae4") # Admin user
        )
        db.add(note)
        db.commit()
        print("Note successfully seeded in database!")
    except Exception as e:
        db.rollback()
        print(f"Failed to seed note: {e}")
    finally:
        db.close()

if __name__ == "__main__":
    run_seed()
